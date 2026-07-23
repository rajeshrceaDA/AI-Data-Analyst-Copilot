from docx import Document
from docx.shared import Pt


def export_resume_to_docx(profile, output_path):
    """
    Generate Resume in DOCX format.
    """

    doc = Document()

    # -----------------------------
    # Default Font
    # -----------------------------
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # -----------------------------
    # Header
    # -----------------------------
    name = profile.get("name", "")
    designation = profile.get("designation", "")

    doc.add_heading(name, level=0)

    if designation:
        p = doc.add_paragraph()
        p.add_run(designation).bold = True

    contact = doc.add_paragraph()

    contact.add_run(profile.get("location", ""))

    contact.add_run(" | ")

    contact.add_run(profile.get("phone", ""))

    contact.add_run(" | ")

    contact.add_run(profile.get("email", ""))

    linkedin = profile.get("linkedin", "")
    github = profile.get("github", "")

    links = []

    if linkedin:
        links.append(f"LinkedIn: {linkedin}")

    if github:
        links.append(f"GitHub: {github}")

    if links:
        doc.add_paragraph(" | ".join(links))

    # -----------------------------
    # Summary
    # -----------------------------
    summary = profile.get("career_goal", "")

    if summary:
        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(summary)

    # -----------------------------
    # Skills
    # -----------------------------
    skills = profile.get("skills", [])

    if skills:

        doc.add_heading("Technical Skills", level=1)

        doc.add_paragraph(" | ".join(skills))

    # -----------------------------
    # Experience
    # -----------------------------
    doc.add_heading("Work Experience", level=1)

    for i in [1, 2]:

        suffix = "" if i == 1 else "_2"

        job = profile.get(f"job_title{suffix}", "")

        if job == "":
            continue

        doc.add_heading(job, level=2)

        doc.add_paragraph(profile.get(f"company_name{suffix}", ""))

        doc.add_paragraph(
            f"{profile.get(f'start_date{suffix}', '')} - {profile.get(f'end_date{suffix}', '')}"
        )

        responsibilities = profile.get(f"responsibilities{suffix}", "")

        if responsibilities:

            doc.add_heading("Key Responsibilities", level=3)

            for item in responsibilities.split("•"):

                item = item.strip()

                if item:
                    doc.add_paragraph(item, style="List Bullet")

        achievements = profile.get(f"achievements{suffix}", "")

        if achievements:

            doc.add_heading("Achievements", level=3)

            for item in achievements.split("•"):

                item = item.strip()

                if item:
                    doc.add_paragraph(item, style="List Bullet")
    # -----------------------------
    # Education
    # -----------------------------
    doc.add_heading("Education", level=1)

    doc.add_paragraph(profile.get("degree", ""))

    doc.add_paragraph(profile.get("college", ""))

    doc.add_paragraph(
        f"{profile.get('passing_year','')} | {profile.get('percentage','')}"
    )

    # -----------------------------
    # Projects
    # -----------------------------
    has_project = False

    for i in [1, 2, 3]:

        name = profile.get(f"project_name_{i}", "")
        tech = profile.get(f"project_tech_{i}", "")
        duration = profile.get(f"project_duration_{i}", "")
        github = profile.get(f"project_github_{i}", "")
        description = profile.get(f"project_description_{i}", "")

        print("-------------------------")
        print(name)
        print(tech)
        print(duration)
        print(github)
        if name == "":
            continue

        if not has_project:
            doc.add_heading("Projects", level=1)
            has_project = True

        doc.add_heading(name, level=2)

        if tech:
            doc.add_paragraph(f"Technologies: {tech}")

        if duration:
            doc.add_paragraph(f"Duration: {duration}")

        if github:
            doc.add_paragraph(f"GitHub: {github}")

        if description:
            doc.add_paragraph(description)

    # -----------------------------
    # Certifications
    # -----------------------------
    certifications = profile.get("certifications", "")

    if certifications:
        doc.add_heading("Certifications", level=1)
        for cert in certifications.split("\n"):

            cert = cert.strip()

            if cert:
                doc.add_paragraph(cert, style="List Bullet")

    # -----------------------------
    # Languages
    # -----------------------------
    languages = profile.get("languages", "")

    if languages:
        doc.add_heading("Languages", level=1)
        for lang in languages.split(","):

            lang = lang.strip()

            if lang:
                doc.add_paragraph(lang, style="List Bullet")

    # -----------------------------
    # Save File
    # -----------------------------
    doc.save(output_path)

    return output_path